"""
The Pool — intake v0.3 (FINAL: full extraction, 100% coverage)

Three ways to read a CV, tried cheapest-first:
  1. PDF with real text    -> pdfplumber (free, instant)
  2. PDF that's a scan/img -> render pages to images -> vision model reads them ($)
  3. Word (.docx) file     -> read the text (with a backup for damaged files, free)

Memory rule: skip a CV once we have GOOD text for it. Anything unresolved gets
retried when we gain a tool that can help. This version's new tool is the vision
model, which rescues the scanned ('suspicious') PDFs.
"""

import os
import io
import sys
import csv
import base64
import hashlib
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import yaml
import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from dotenv import load_dotenv
from anthropic import Anthropic

sys.stdout.reconfigure(encoding="utf-8")

# --- Load config + secret key ---
BASE = Path(__file__).parent
config = yaml.safe_load((BASE / "config.yaml").read_text(encoding="utf-8"))

extracted_dir = BASE / config["paths"]["extracted"]
ledger_path   = BASE / config["paths"]["processed_ledger"]
results_path  = BASE / config["paths"]["results_report"]
threshold     = config["settings"]["suspicious_char_threshold"]
extracted_dir.mkdir(exist_ok=True)

load_dotenv()                                   # load .env into the environment
if not os.getenv("ANTHROPIC_API_KEY"):
    print("WARNING: no ANTHROPIC_API_KEY in .env — the vision rescue will fail.\n")
client = Anthropic()                            # reads ANTHROPIC_API_KEY automatically

# --- Vision rescue settings ---
VISION_MODEL = "claude-sonnet-5"                # good at reading text-heavy images
RENDER_SCALE = 2.0                              # page image resolution (2.0 is a solid default)
MAX_PAGES    = 10                               # safety cap so a giant scan can't run up the bill
VISION_PROMPT = (
    "This is a scanned CV / resume, given as one or more page images. "
    "Transcribe ALL the text exactly as it appears, keeping the reading order "
    "and structure (headings, sections, bullet points, dates). "
    "Include every language present (for example both Arabic and English). "
    "Do NOT translate, summarize, correct, or add any commentary — "
    "output only the transcribed text of the CV."
)

# Retry a CV whose status is one of these — we now have a tool that can help.
RETRY_STATES = {"suspicious", "failed"}

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
LEDGER_COLUMNS = ["hash", "source", "file", "chars", "status", "note"]
REPORT_COLUMNS = ["source", "file", "chars", "status", "note"]


def file_fingerprint(path):
    """Unique fingerprint of a file's contents (same contents -> same fingerprint)."""
    return hashlib.sha256(path.read_bytes()).hexdigest()


def extract_pdf_text(path):
    """Path 1: pull text out of a normal PDF with pdfplumber."""
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def _docx_text_from_zip(path):
    """Backup Word reader: a .docx is a zip; the words live in word/document.xml.
    Read just that part and ignore any broken images."""
    with zipfile.ZipFile(path) as z:
        xml_bytes = z.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    lines = []
    for para in root.iter(f"{W_NS}p"):
        pieces = [node.text for node in para.iter(f"{W_NS}t") if node.text]
        line = "".join(pieces).strip()
        if line:
            lines.append(line)
    return "\n".join(lines).strip()


def extract_docx_text(path):
    """Path 3: read a Word file. Normal path (gets tables); if the zip is damaged,
    fall back to reading the text XML directly."""
    try:
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(p for p in parts if p.strip()).strip()
    except Exception:
        return _docx_text_from_zip(path)


def pdf_pages_as_png(path):
    """Render each PDF page to PNG bytes, so the vision model can 'see' them."""
    images = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        for i in range(min(len(pdf), MAX_PAGES)):
            bitmap = pdf[i].render(scale=RENDER_SCALE)
            buffer = io.BytesIO()
            bitmap.to_pil().save(buffer, format="PNG")
            images.append(buffer.getvalue())
    finally:
        pdf.close()
    return images


def read_with_vision(png_images):
    """Path 2: send page images to the vision model; get transcribed text back.
    Returns (text, input_tokens, output_tokens)."""
    content = []
    for img in png_images:
        content.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/png",
                "data": base64.standard_b64encode(img).decode("utf-8"),
            },
        })
    content.append({"type": "text", "text": VISION_PROMPT})   # images first, then the ask

    message = client.messages.create(
        model=VISION_MODEL,
        max_tokens=8192,
        messages=[{"role": "user", "content": content}],
    )
    text = "".join(b.text for b in message.content if b.type == "text").strip()
    return text, message.usage.input_tokens, message.usage.output_tokens


# --- Load the ledger (memory) ---
seen = {}
if ledger_path.exists():
    with ledger_path.open("r", newline="", encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            seen[row["hash"]] = row
print(f"Ledger loaded: {len(seen)} CVs already in the pool.\n")

# --- The loop ---
handled = {}
left_alone = 0
duplicates = 0
vision_calls = 0
input_tokens = 0
output_tokens = 0


def bump(status):
    handled[status] = handled.get(status, 0) + 1


for source, folder in config["sources"].items():
    source_dir = BASE / folder
    if not source_dir.exists():
        continue

    for file in [f for f in source_dir.rglob("*") if f.is_file()]:
        rel = file.relative_to(source_dir)
        rel_str = str(rel)

        try:
            fp = file_fingerprint(file)
        except Exception as error:
            bump("failed")
            print(f"[{source}] {rel}  ->  FAILED (couldn't read file: {error})")
            continue

        if fp in seen:
            prev = seen[fp]
            same_file = (prev["source"] == source and prev["file"] == rel_str)
            if not same_file:
                duplicates += 1
                print(f"[{source}] {rel}  ->  DUPLICATE of [{prev['source']}] {prev['file']} (skipped)")
                continue
            if prev["status"] not in RETRY_STATES:
                left_alone += 1
                continue
            print(f"[{source}] {rel}  ->  rescuing (was '{prev['status']}')...")

        # --- Process (new file, or a retry) ---
        suffix = file.suffix.lower()
        text, status, note = "", "", ""

        if suffix == ".pdf":
            try:
                text = extract_pdf_text(file)
            except Exception as error:
                text = ""
                print(f"    pdfplumber couldn't read it ({error}) — trying vision")

            if len(text) >= threshold:
                status, note = "clean", "pdfplumber"
                print(f"[{source}] {rel}  ->  clean ({len(text)} chars)")
            else:
                # scan / image PDF -> vision rescue
                try:
                    images = pdf_pages_as_png(file)
                    if not images:
                        status, note = "vlm_empty", "no pages to render"
                    else:
                        text, itok, otok = read_with_vision(images)
                        vision_calls += 1
                        input_tokens += itok
                        output_tokens += otok
                        if len(text) >= threshold:
                            status, note = "clean", "vision-rescued"
                        else:
                            status, note = "vlm_empty", "vision found little/no text"
                    print(f"[{source}] {rel}  ->  {status} ({len(text)} chars, via vision)")
                except Exception as error:
                    text, status, note = "", "failed", f"vision error: {error}"
                    print(f"[{source}] {rel}  ->  FAILED vision ({error})")

        elif suffix == ".docx":
            try:
                text = extract_docx_text(file)
                if len(text) >= threshold:
                    status, note = "clean", "docx"
                else:
                    status, note = "docx_empty", "Word file had little/no text"
            except Exception as error:
                status, note = "failed", f"docx error: {error}"
            print(f"[{source}] {rel}  ->  {status} ({len(text)} chars, docx)")

        else:
            status, note = "unsupported", f"unsupported type ({file.suffix})"
            print(f"[{source}] {rel}  ->  unsupported ({file.suffix})")

        out_name = f"{source}__" + "__".join(rel.with_suffix("").parts) + ".txt"
        (extracted_dir / out_name).write_text(text, encoding="utf-8")

        bump(status)
        seen[fp] = {"hash": fp, "source": source, "file": rel_str,
                    "chars": len(text), "status": status, "note": note}

# --- Save ledger + results ---
with ledger_path.open("w", newline="", encoding="utf-8-sig") as f:
    writer = csv.DictWriter(f, fieldnames=LEDGER_COLUMNS)
    writer.writeheader()
    writer.writerows(seen.values())

with results_path.open("w", newline="", encoding="utf-8-sig") as f:
    writer = csv.DictWriter(f, fieldnames=REPORT_COLUMNS)
    writer.writeheader()
    for row in seen.values():
        writer.writerow({k: row[k] for k in REPORT_COLUMNS})

# --- Summary ---
print("\n" + "=" * 50)
print("THIS RUN")
print(f"  handled this run : {sum(handled.values())}")
for status in sorted(handled):
    print(f"      {status:<16}: {handled[status]}")
print(f"  left untouched   : {left_alone}")
print(f"  duplicates       : {duplicates}")
if vision_calls:
    print(f"  vision calls     : {vision_calls}  "
          f"({input_tokens:,} in + {output_tokens:,} out tokens)")
print("-" * 50)
pool = {}
for row in seen.values():
    pool[row["status"]] = pool.get(row["status"], 0) + 1
total = sum(pool.values())
clean = pool.get("clean", 0)
print(f"POOL NOW: {total} CVs total, {clean} readable ({clean / total * 100:.0f}%)")
for status in sorted(pool):
    print(f"      {status:<16}: {pool[status]}")
print("=" * 50)